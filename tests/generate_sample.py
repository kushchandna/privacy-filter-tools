#!/usr/bin/env python3
"""
Generate tests/output/sample.docx with synthetic PII for smoke testing.

Covers all 8 opf detection categories, plus two regression-tracker items:
  - ALL-CAPS name in a heading  (known model recall gap for upper-case names)
  - Literal "&" in a heading    (regression: html.unescape must strip &amp; leaks)

Usage:
    python tests/generate_sample.py                       # default path
    python tests/generate_sample.py path/to/output.docx  # custom path
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Synthetic PII constants — clearly fake, never real personal data.
# Exported so smoke_test.py can import and check they are absent from output.
# ---------------------------------------------------------------------------
SYNTHETIC_NAME = "Alice Redactorina Johnson"
SYNTHETIC_EMAIL = "alice.r.johnson@synth-example.invalid"
SYNTHETIC_PHONE = "+1 (555) 867-5309"
SYNTHETIC_ADDRESS = "742 Evergreen Terrace, Springfield, IL 62701"
SYNTHETIC_DOB = "1990-03-15"
SYNTHETIC_USERNAME = "@alice_redactorina_90"
SYNTHETIC_ID_NUM = "SSN: 523-45-6789"
SYNTHETIC_CREDENTIAL = "API_KEY=sk-synth1234567890abcdef"

# Regression tracker: model sometimes misses names written entirely in caps.
CAPS_NAME = "ALICE REDACTORINA JOHNSON"

_DEFAULT_OUTPUT = Path(__file__).parent / "output" / "sample.docx"


def generate(output_path: Path) -> None:
    """Write a sample .docx to *output_path*, creating parent dirs as needed."""
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        sys.exit(f"generate_sample: python-docx is required — pip install python-docx ({exc})")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Regression tracker: ALL-CAPS name (model recall gap)
    doc.add_heading(f"CONTRACT FOR {CAPS_NAME}", level=1)

    # Regression tracker: literal & (must not appear as &amp; in converted Markdown)
    doc.add_heading("Terms & Conditions", level=2)

    # Name + date of birth (categories 1 & 5)
    doc.add_paragraph(
        f"This agreement is entered into by {SYNTHETIC_NAME} "
        f"(date of birth: {SYNTHETIC_DOB})."
    )

    # Email + phone (categories 2 & 3)
    doc.add_paragraph(
        f"Contact information: {SYNTHETIC_EMAIL}, phone {SYNTHETIC_PHONE}."
    )

    # Address (category 4)
    doc.add_paragraph(f"Mailing address: {SYNTHETIC_ADDRESS}.")

    # Government ID / ID number (category 6)
    doc.add_paragraph(f"Government ID: {SYNTHETIC_ID_NUM}.")

    # Username / handle (category 7)
    doc.add_paragraph(f"Online handle: {SYNTHETIC_USERNAME}.")

    # Credential / API key (category 8)
    doc.add_paragraph(f"Credentials on file: {SYNTHETIC_CREDENTIAL}.")

    doc.save(output_path)
    print(f"Generated {output_path}")


def _main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "output",
        nargs="?",
        default=str(_DEFAULT_OUTPUT),
        help=f"Output .docx path (default: {_DEFAULT_OUTPUT})",
    )
    args = parser.parse_args()
    generate(Path(args.output))


if __name__ == "__main__":
    _main()
